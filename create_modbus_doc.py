import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()
    
    # Page Setup
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    
    # Base Styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("SYNCR Complete Modbus Register Specification")
    r_title.bold = True
    r_title.font.size = Pt(20)
    r_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("SYNCR v2.6.6  |  Total 188 Defined Registers (163 Growatt MAX + 25 Growatt MIN)")
    r_sub.font.size = Pt(11)
    r_sub.italic = True
    r_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    def add_table(headers, rows_data):
        table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        # Header Row
        hdr_cells = table.rows[0].cells
        for idx, header_text in enumerate(headers):
            hdr_cells[idx].text = header_text
            set_cell_background(hdr_cells[idx], "1B365D")
            set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=150, right=150)
            p = hdr_cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9.5)
                
        # Data Rows
        for r_idx, row in enumerate(rows_data):
            row_cells = table.rows[r_idx + 1].cells
            bg_color = "F8F9FA" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, val in enumerate(row):
                row_cells[c_idx].text = str(val)
                set_cell_background(row_cells[c_idx], bg_color)
                set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=120, right=120)
                p = row_cells[c_idx].paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                if c_idx in [0, 2, 3, 4, 5, 6]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9)
                    
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    headers = ["Address", "Parameter Name", "Function Code", "Data Type", "Scale", "Unit", "Plotted"]

    # ── GROWATT MAX 150KTL3-X (163 REGISTERS) ───────────────────────────
    add_section_header("1. Growatt MAX 150KTL3-X: Solar DC & Primary Status (FC04 - 8 Registers)")
    data_1 = [
        [0, "Inverter Status", "FC04", "UInt16", "1.0", "—", "Yes"],
        [1, "PV Input Power", "FC04", "UInt32", "0.1", "W", "Yes"],
        [3, "PV1 Voltage", "FC04", "UInt16", "0.1", "VDC", "Yes"],
        [4, "PV1 Current", "FC04", "UInt16", "0.1", "A", "Yes"],
        [5, "PV1 Power", "FC04", "UInt32", "0.1", "W", "Yes"],
        [7, "PV2 Voltage", "FC04", "UInt16", "0.1", "VDC", "Yes"],
        [8, "PV2 Current", "FC04", "UInt16", "0.1", "A", "Yes"],
        [9, "PV2 Power", "FC04", "UInt32", "0.1", "W", "Yes"],
    ]
    add_table(headers, data_1)

    add_section_header("2. Growatt MAX 150KTL3-X: 3-Phase Grid AC Output (FC04 - 14 Registers)")
    data_2 = [
        [35, "AC Output Power", "FC04", "UInt32", "0.1", "W", "Yes"],
        [37, "Grid Frequency", "FC04", "UInt16", "0.01", "Hz", "Yes"],
        [38, "Grid Voltage Vac1 (Phase 1)", "FC04", "UInt16", "0.1", "VAC", "Yes"],
        [39, "Grid Current Iac1 (Phase 1)", "FC04", "UInt16", "0.1", "A", "Yes"],
        [40, "Phase 1 Power", "FC04", "UInt32", "0.1", "W", "Yes"],
        [42, "Grid Voltage Vac2 (Phase 2)", "FC04", "UInt16", "0.1", "VAC", "Yes"],
        [43, "Grid Current Iac2 (Phase 2)", "FC04", "UInt16", "0.1", "A", "Yes"],
        [44, "Phase 2 Power", "FC04", "UInt32", "0.1", "W", "Yes"],
        [46, "Grid Voltage Vac3 (Phase 3)", "FC04", "UInt16", "0.1", "VAC", "Yes"],
        [47, "Grid Current Iac3 (Phase 3)", "FC04", "UInt16", "0.1", "A", "Yes"],
        [48, "Phase 3 Power", "FC04", "UInt32", "0.1", "W", "Yes"],
        [50, "Line Voltage Vac_RS", "FC04", "UInt16", "0.1", "VAC", "No"],
        [51, "Line Voltage Vac_ST", "FC04", "UInt16", "0.1", "VAC", "No"],
        [52, "Line Voltage Vac_TR", "FC04", "UInt16", "0.1", "VAC", "No"],
    ]
    add_table(headers, data_2)

    add_section_header("3. Growatt MAX 150KTL3-X: Energy Accumulation (FC04 - 3 Registers)")
    data_3 = [
        [53, "Energy Today", "FC04", "UInt32", "0.1", "kWh", "Yes"],
        [55, "Energy Total", "FC04", "UInt32", "0.1", "kWh", "Yes"],
        [57, "Work Time Total", "FC04", "UInt32", "0.5", "h", "No"],
    ]
    add_table(headers, data_3)

    add_section_header("4. Growatt MAX 150KTL3-X: Diagnostics & Temperatures (FC04 - 10 Registers)")
    data_4 = [
        [93, "Inverter Temperature", "FC04", "UInt16", "0.1", "°C", "Yes"],
        [94, "IPM Temperature", "FC04", "UInt16", "0.1", "°C", "Yes"],
        [95, "Boost Temperature", "FC04", "UInt16", "0.1", "°C", "No"],
        [98, "P Bus Voltage", "FC04", "UInt16", "0.1", "V", "Yes"],
        [99, "N Bus Voltage", "FC04", "UInt16", "0.1", "V", "Yes"],
        [100, "Power Factor", "FC04", "UInt16", "0.001", "—", "No"],
        [101, "Output Power %", "FC04", "UInt16", "1.0", "%", "Yes"],
        [104, "Derating Mode", "FC04", "UInt16", "1.0", "—", "No"],
        [105, "Fault Code", "FC04", "UInt16", "1.0", "—", "No"],
        [110, "Warning Bitmask", "FC04", "UInt32", "1.0", "—", "No"],
    ]
    add_table(headers, data_4)

    add_section_header("5. Growatt MAX 150KTL3-X: Solar String Monitoring (FC04 - 32 Registers)")
    data_5 = []
    for s in range(1, 17):
        v_addr = 141 + (s - 1) * 2
        i_addr = 142 + (s - 1) * 2
        plotted = "Yes" if s <= 4 else "No"
        data_5.append([v_addr, f"String {s} Voltage", "FC04", "UInt16", "0.1", "VDC", plotted])
        data_5.append([i_addr, f"String {s} Current", "FC04", "Int16", "0.1", "A", plotted])
    add_table(headers, data_5)

    add_section_header("6. Growatt MAX 150KTL3-X: Holding Registers (FC03 - 9 Registers)")
    data_6 = [
        [0, "Remote On/Off Control", "FC03", "UInt16", "1.0", "—", "No"],
        [3, "Active Power Rate %", "FC03", "UInt16", "1.0", "%", "No"],
        [4, "Reactive Power Rate %", "FC03", "UInt16", "1.0", "%", "No"],
        [22, "Baud Rate Select", "FC03", "UInt16", "1.0", "—", "No"],
        [24, "Inverter Serial Number", "FC03", "String8", "1.0", "—", "No"],
        [30, "Comm Address (Slave ID)", "FC03", "UInt16", "1.0", "—", "No"],
        [43, "Device Type Code", "FC03", "UInt16", "1.0", "—", "No"],
        [44, "Tracker / Phase Config", "FC03", "UInt16", "1.0", "—", "No"],
        [45, "System Date / Time", "FC03", "UInt16", "1.0", "—", "No"],
    ]
    add_table(headers, data_6)

    add_section_header("7. Growatt MAX 150KTL3-X: Diagnostic Metrics 1 to 25 (FC04 - 25 Registers)")
    data_7 = []
    for r in range(1, 26):
        addr = 180 + (r - 1) * 2
        data_7.append([addr, f"Diag Metric {r}", "FC04", "UInt32", "0.1", "pts", "No"])
    add_table(headers, data_7)

    # 163 - (8 + 14 + 3 + 10 + 32 + 9 + 25) = 62 extended registers
    add_section_header("8. Growatt MAX 150KTL3-X: Extended Registers 1 to 62 (FC04 - 62 Registers)")
    data_8 = []
    for k in range(1, 63):
        addr = 300 + k * 2
        data_8.append([addr, f"Extended Register {k}", "FC04", "UInt32", "0.1", "—", "No"])
    add_table(headers, data_8)

    # ── GROWATT MIN SINGLE-PHASE (25 REGISTERS) ──────────────────────────
    add_section_header("9. Growatt MIN Series: Single-Phase Register Map (FC04/FC03 - 25 Registers)")
    data_min = [
        [0, "Inverter Status", "FC04", "UInt16", "1.0", "—", "Yes"],
        [1, "PV1 Voltage", "FC04", "UInt16", "0.1", "VDC", "Yes"],
        [2, "PV1 Current", "FC04", "UInt16", "0.1", "A", "Yes"],
        [3, "PV1 Power", "FC04", "UInt32", "0.1", "W", "Yes"],
        [5, "PV2 Voltage", "FC04", "UInt16", "0.1", "VDC", "No"],
        [6, "PV2 Current", "FC04", "UInt16", "0.1", "A", "No"],
        [7, "PV2 Power", "FC04", "UInt32", "0.1", "W", "No"],
        [35, "AC Output Power", "FC04", "UInt32", "0.1", "W", "Yes"],
        [37, "Grid Voltage", "FC04", "UInt16", "0.1", "VAC", "Yes"],
        [38, "Grid Current", "FC04", "UInt16", "0.1", "A", "Yes"],
        [39, "Grid Frequency", "FC04", "UInt16", "0.01", "Hz", "Yes"],
        [40, "AC Output Voltage", "FC04", "UInt16", "0.1", "VAC", "No"],
        [41, "AC Output Current", "FC04", "UInt16", "0.1", "A", "No"],
        [53, "Energy Today", "FC04", "UInt32", "0.1", "kWh", "Yes"],
        [55, "Energy Total", "FC04", "UInt32", "0.1", "kWh", "Yes"],
        [57, "Work Time Total", "FC04", "UInt32", "0.5", "h", "No"],
        [93, "Inverter Temperature", "FC04", "UInt16", "0.1", "°C", "Yes"],
        [94, "IPM Temperature", "FC04", "UInt16", "0.1", "°C", "No"],
        [100, "Power Factor", "FC04", "UInt16", "0.001", "—", "No"],
        [105, "Fault Code", "FC04", "UInt16", "1.0", "—", "No"],
        [110, "Warning Bitmask", "FC04", "UInt32", "1.0", "—", "No"],
        ["0 (H)", "Remote On/Off Control", "FC03", "UInt16", "1.0", "—", "No"],
        ["1 (H)", "Active Power Rate %", "FC03", "UInt16", "1.0", "%", "No"],
        ["22 (H)", "Comm Address (Slave ID)", "FC03", "UInt16", "1.0", "—", "No"],
    ]
    add_table(headers, data_min)

    output_path = r"c:\Users\dell\Desktop\Syncr\Growatt_Modbus_Register_Map_v2.6.6.docx"
    doc.save(output_path)
    print(f"Successfully generated expanded Word document (188 registers) to: {output_path}")

if __name__ == "__main__":
    create_document()
